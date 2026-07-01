#!/usr/bin/env python3
"""
generate_reports.py
Build aligned HTML and DOCX reports from eda-field-level output artifacts.
"""

import os
from pathlib import Path
from _discover_growers import resolve_output_dir, resolve_data_root

# ---------------------------------------------------------------------------
# Report content  (single source of truth for both formats)
# ---------------------------------------------------------------------------

GROWER_TABLE = [
    ("Grower", "State", "Fields", "Farm"),
    ("ia-grower", "Iowa", "10", "ia-grower-iowa"),
    ("il-grower", "Illinois", "10", "il-grower-illinois"),
    ("ne-grower", "Nebraska", "10", "ne-grower-nebraska"),
    ("Total", "", "30 fields", "3 farms"),
]

LAYERS_TABLE = [
    ("Layer", "Source", "Period", "Format"),
    (
        "Field boundaries",
        "OSM / farm boundary files",
        "Static",
        "GeoJSON (EPSG:4326, projected to EPSG:5070 for area)",
    ),
    (
        "Cropland Data Layer (CDL)",
        "USDA NASS CDL",
        "2021\u20132025",
        "CSV composition tables (crop code, pixel count per field-year)",
    ),
    (
        "Weather",
        "NASA POWER (PRECTOTCORR, T2M)",
        "2021\u20132025",
        "CSV daily records",
    ),
]

COMPARISON_LEVELS = [
    (
        "Field",
        "Individual field size (hectares) and perimeter compared within and "
        "across growers; per-field dominant crop tracked across years in rotation heatmaps.",
    ),
    (
        "Field-year",
        "Dominant CDL crop identity per field per year (the unit used in rotation "
        "heatmaps and yearly composition bars).",
    ),
    (
        "Grower",
        "Per-grower aggregations \u2014 total acreage, crop composition by year, mean "
        "daily precipitation by season, and annual temperature/precipitation summaries.",
    ),
    (
        "Across-grower",
        "Cross-grower charts that overlay or juxtapose all three growers \u2014 geospatial "
        "centroid map, field-size histograms, cumulative area bars, field-count vs. acreage "
        "paired bars, CDL composition flow, and annual precipitation/temperature comparisons.",
    ),
]

CATEGORIES = [
    # ------------------------------------------------------------------ Field Boundaries
    {
        "id": "field-boundaries",
        "title": "Field Boundaries",
        "analysis": (
            "Field size and shape directly impact farm operations \u2014 larger, regular fields "
            "allow wider equipment passes and fewer turns, while smaller or irregular fields "
            "require more precise handling and increase edge effects. Comparing geometry across "
            "growers reveals differences in operational scale, landscape context, and potential "
            "management intensity."
        ),
        "csv_outputs": [
            {
                "filename": "cross_grower_field_boundary_metrics.csv",
                "desc": "Per-field geometry: grower, field_id, area_ha, perimeter_m for all 30 fields.",
            },
            {
                "filename": "cross_grower_field_count_acreage_summary.csv",
                "desc": "Per-grower summary: field count and total area (hectares) for each of the 3 growers.",
            },
            {
                "filename": "cross_grower_field_centroids.csv",
                "desc": "Field centroid coordinates (lon, lat) plus area_acres and original GeoJSON attributes for all 30 fields.",
            },
        ],
        "png_outputs": [
            {
                "filename": "cross_grower_field_centroid_map.png",
                "desc": (
                    "30 field centroids on a state-outline basemap with coastline-clipped state fills. "
                    "Marker size scales with field area; color distinguishes growers. Lake water bodies "
                    "appear as blue overlays. A 2\u00b0 \u00d7 1\u00b0 lat/lon graticule provides "
                    "spatial reference."
                ),
                "analysis": (
                    "Places the geometry statistics into a geographic context. Proximity, clustering, "
                    "and spatial separation between growers\u2019 fields become visible at a glance \u2014 "
                    "do all fields cluster in one county, or are they spread across a region?"
                ),
            },
            {
                "filename": "cross_grower_field_area_cumulative_stacked_bar.png",
                "desc": (
                    "One bar per grower; each segment is one field stacked largest to smallest. "
                    "Bar height shows total farm area; the number of segments reveals how many fields "
                    "compose it."
                ),
                "analysis": (
                    "Reveals how each grower\u2019s total acreage is distributed \u2014 is it "
                    "concentrated in a few large fields or spread across many smaller ones? A grower "
                    "with a tall bar and many small segments operates many small fields; a tall bar "
                    "with few large segments indicates large-field agriculture."
                ),
            },
            {
                "filename": "cross_grower_field_area_histogram.png",
                "desc": (
                    "Faceted panels with shared bin edges. Each panel shows one grower\u2019s "
                    "field-size distribution with median area marked as a dashed vertical line."
                ),
                "analysis": (
                    "Shows the distribution shape and spread of field sizes within each grower. "
                    "A narrow peak around a high median indicates uniformly large fields; a wide, "
                    "right-skewed distribution suggests a mix of field sizes with some very large outliers."
                ),
            },
            {
                "filename": "cross_grower_field_count_vs_acreage.png",
                "desc": (
                    "Dual-axis bar chart: field count (left axis, opaque bars) vs. total area "
                    "(right axis, translucent bars) per grower."
                ),
                "analysis": (
                    "Highlights growers with many small fields vs. fewer large fields. When field "
                    "count and total area bars are disproportionate, it signals that a grower\u2019s "
                    "average field size differs markedly from the others \u2014 a key insight for "
                    "comparing operational scale."
                ),
            },
        ],
    },
    # ------------------------------------------------------------------ CDL Crop Data
    {
        "id": "cdl-crop-data",
        "title": "CDL Crop Data",
        "analysis": (
            "Crop rotation is one of the most powerful management tools available \u2014 it breaks "
            "pest cycles, manages soil fertility, and spreads economic risk. The USDA Cropland Data "
            "Layer (CDL) provides annual 30 m resolution crop classifications that let us trace which "
            "crops were planted where, year after year. This section examines crop composition, "
            "rotation patterns, and diversity across growers and fields."
        ),
        "csv_outputs": [
            {
                "filename": "cross_grower_cdl_composition.csv",
                "desc": "Full CDL composition: field_id, year, crop_code, crop_name, pixel_count, pct per field-year-crop.",
            },
            {
                "filename": "cross_grower_cdl_crop_totals_by_year.csv",
                "desc": "Aggregated crop totals: grower, year, crop_name, pixel_count summed across fields.",
            },
        ],
        "png_outputs": [
            {
                "filename": "cross_grower_cdl_crop_composition_flow.png",
                "desc": (
                    "One panel per grower: stacked bars per year with translucent connecting bands "
                    "between successive years. Each segment and band is colored by crop type."
                ),
                "analysis": (
                    "Makes year-to-year crop transitions visible. Wide bands of the same color indicate "
                    "stable rotations (e.g., corn every year); changing colors reveal substitutions "
                    "(e.g., corn to soybeans). The flow bands emphasize transitions that a simple bar "
                    "chart alone would hide."
                ),
            },
            {
                "template": "{grower}_cdl_crop_composition_by_year.png",
                "actual_files": ["ia-grower_cdl_crop_composition_by_year.png",
                                 "il-grower_cdl_crop_composition_by_year.png",
                                 "ne-grower_cdl_crop_composition_by_year.png"],
                "desc": (
                    "One stacked bar per year per grower, each segment proportional to that crop\u2019s "
                    "share of total CDL pixels."
                ),
                "analysis": (
                    "Reveals each grower\u2019s aggregate crop mix and how it shifts over the 5-year "
                    "window. A grower shifting from corn-dominated to more diverse plantings shows "
                    "visible composition changes year over year."
                ),
            },
            {
                "template": "{grower}_cdl_crop_rotation_heatmap.png",
                "actual_files": ["ia-grower_cdl_crop_rotation_heatmap.png",
                                 "il-grower_cdl_crop_rotation_heatmap.png",
                                 "ne-grower_cdl_crop_rotation_heatmap.png"],
                "desc": (
                    "Per-grower heatmap: rows = field IDs, columns = years, cell text = dominant crop "
                    "name. Colored by crop type."
                ),
                "analysis": (
                    "Makes individual field crop sequences legible at a glance. A classic corn\u2013soy "
                    "rotation appears as alternating yellow and green rows. A field that stays in corn "
                    "or shifts to wheat stands out immediately as a departure from the dominant rotation "
                    "pattern."
                ),
            },
        ],
    },
    # ------------------------------------------------------------------ Weather
    {
        "id": "weather",
        "title": "Weather",
        "analysis": (
            "Weather, especially precipitation during the growing season, is the dominant "
            "uncontrollable variable in crop production. Understanding the range, timing, and "
            "variability of rainfall and temperature across years helps contextualize yield outcomes "
            "and gauge production risk. NASA POWER provides global gridded daily weather estimates "
            "for each field location."
        ),
        "csv_outputs": [
            {
                "filename": "cross_grower_weather_annual_precipitation.csv",
                "desc": "Grower-year total precipitation (mean and std across fields within each grower).",
            },
            {
                "filename": "cross_grower_weather_annual_summary.csv",
                "desc": "Grower-year total_precip and mean_temp for dual-axis comparison.",
            },
        ],
        "png_outputs": [
            {
                "filename": "cross_grower_weather_average_cumulative_annual_precip.png",
                "desc": (
                    "Grouped bar chart with per-field standard-deviation error bars. Each grower has "
                    "one bar per year, grouped by year across growers."
                ),
                "analysis": (
                    "Compares total annual precipitation across growers and years. Taller bars indicate "
                    "wetter years; error bars show within-grower field-to-field variability \u2014 "
                    "larger error bars suggest spatially heterogeneous rainfall across a grower\u2019s fields."
                ),
            },
            {
                "filename": "cross_grower_weather_temp_precip_dual_axis.png",
                "desc": (
                    "Precipitation bars (left axis) overlaid with mean annual temperature line "
                    "(right axis) per grower per year."
                ),
                "analysis": (
                    "A hot, dry year jumps out when a low precip bar sits beneath a high temperature "
                    "point. Cool, wet years are equally visible. This dual-axis view makes it easy to "
                    "spot correlated climate extremes across the three locations."
                ),
            },
            {
                "template": "{grower}_weather_mean_daily_precip_by_season.png",
                "actual_files": ["ia-grower_weather_mean_daily_precip_by_season.png",
                                 "il-grower_weather_mean_daily_precip_by_season.png",
                                 "ne-grower_weather_mean_daily_precip_by_season.png"],
                "docx_img_width": 3.5,
                "desc": (
                    "One panel per grower, faceted by year. Bars colored by season "
                    "(blue = winter, green = spring, orange = summer, red = fall)."
                ),
                "analysis": (
                    "Shows when precipitation fell within each year. Are spring rains arriving earlier "
                    "or later? Are summers consistently dry? The seasonal breakdown reveals timing "
                    "patterns that total annual precipitation alone masks."
                ),
            },
        ],
    },
]

KEY_OBSERVATIONS = [
    (
        "Field counts are uniform",
        "All three growers have exactly 10 fields each, making per-grower comparisons "
        "straightforward without count normalization.",
    ),
    (
        "Corn\u2013soybean rotation dominates",
        "Iowa and Illinois fields are overwhelmingly corn and soybeans (typical US Midwest "
        "pattern). Nebraska shows greater crop diversity including winter wheat and fallow/idle acreage.",
    ),
    (
        "Precipitation gradient",
        "Illinois typically receives the highest annual precipitation, followed by Iowa, with "
        "Nebraska the driest. The seasonal precipitation plots show distinct spring/summer peaks "
        "in all three locations.",
    ),
    (
        "Field size distribution",
        "The faceted histograms and cumulative bars reveal whether a grower\u2019s acreage is "
        "concentrated in a few large fields or spread across many smaller ones. Median field sizes "
        "differ across states.",
    ),
]

LIMITATIONS = [
    (
        "Limited temporal window",
        "CDL and weather data span 2021\u20132025 (5 years). Longer-term trends cannot be assessed.",
    ),
    (
        "Small grower sample",
        "Three growers (IA, IL, NE) covering only three Midwestern states. Results are not "
        "representative of US agriculture broadly.",
    ),
    (
        "CDL resolution",
        "30 m pixels may miss small-field or mixed cropping patterns.",
    ),
    (
        "No soil data",
        "Soil properties (texture, OM, CEC, pH) were not included. The field-level workflow "
        "explicitly excludes soil analysis.",
    ),
    (
        "No imagery data",
        "Vegetation indices (NDVI, etc.) from Sentinel-2 or Landsat were not part of this analysis.",
    ),
    (
        "Single farm per grower",
        "Each grower operates only one farm in this dataset. Grower-level comparisons are "
        "effectively farm-level comparisons.",
    ),
    (
        "Limited field sample",
        "The data pipeline was configured to pull 10 fields per grower. Growers may operate more "
        "fields in practice. These 10 were discovered dynamically; results reflect only the sampled subset.",
    ),
]

GROWER_SLUGS = ["ia-grower", "il-grower", "ne-grower"]


def resolve_grower_slugs(output_dir: Path) -> list[str]:
    """Discover grower slugs from existing per-grower PNGs in the output dir."""
    slugs = set()
    for p in output_dir.glob("*.png"):
        parts = p.stem.split("_", 1)
        if parts[0] in GROWER_SLUGS:
            slugs.add(parts[0])
    return sorted(slugs) if slugs else GROWER_SLUGS


# ---------------------------------------------------------------------------
# HTML generator
# ---------------------------------------------------------------------------

def _build_html(output_dir: Path, grower_slugs: list[str]) -> str:
    lines = []
    def w(s: str = ""):
        lines.append(s)
    def esc(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    w("<!DOCTYPE html>")
    w('<html lang="en">')
    w("<head>")
    w('<meta charset="UTF-8">')
    w('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
    w("<title>EDA Field-Level Analysis Report</title>")
    w("<style>")
    w("body {")
    w('  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;')
    w("  max-width: 960px;")
    w("  margin: 0 auto;")
    w("  padding: 2em 1em;")
    w("  line-height: 1.6;")
    w("  color: #222;")
    w("  background: #fafafa;")
    w("}")
    w("h1 { border-bottom: 2px solid #1f77b4; padding-bottom: 0.3em; }")
    w("h2 { color: #1f77b4; margin-top: 2em; }")
    w("h3 { margin-top: 1.5em; color: #444; }")
    w("h4 { margin-top: 1.2em; color: #666; }")
    w("p, li { max-width: 720px; }")
    w("table { border-collapse: collapse; margin: 1em 0; }")
    w("th, td { border: 1px solid #ccc; padding: 0.4em 0.8em; text-align: left; }")
    w("th { background: #1f77b4; color: #fff; }")
    w("tr:nth-child(even) { background: #f0f0f0; }")
    w(".gallery { display: flex; flex-wrap: wrap; gap: 1em; margin: 1.5em 0; }")
    w(".figure { background: #fff; border: 1px solid #ddd; border-radius: 6px; padding: 0.8em; flex: 1 1 420px; margin-bottom: 1.2em; }")
    w(".figure .filename { font-size: 0.82em; font-family: monospace; color: #333; margin-bottom: 0.3em; }")
    w(".figure .desc { font-size: 0.88em; margin-top: 0.2em; color: #555; }")
    w(".figure .analysis { font-size: 0.88em; margin-top: 0.3em; margin-bottom: 0.5em; color: #b85b14; font-style: italic; border-left: 3px solid #b85b14; padding-left: 0.6em; }")
    w(".figure img { max-width: 100%; height: auto; display: block; }")
    w(".category-analysis { background: #e8f4f8; border-left: 4px solid #1f77b4; padding: 0.6em 1em; margin: 1em 0; border-radius: 4px; }")
    w(".limitation { background: #fffbe6; border-left: 4px solid #f0c040; padding: 0.6em 1em; margin: 1em 0; border-radius: 4px; }")
    w(".info { background: #e8f4f8; border-left: 4px solid #1f77b4; padding: 0.6em 1em; margin: 1em 0; border-radius: 4px; }")
    w("code { background: #eee; padding: 0.1em 0.3em; border-radius: 3px; font-size: 0.9em; }")
    w("</style>")
    w("</head>")
    w("<body>")

    w('<h1>EDA Field-Level Analysis Report</h1>')
    w('<p><strong>Generated:</strong> 2026-07-01 &nbsp;|&nbsp; <strong>Subskill:</strong> <code>my-farm-advisor/eda/eda-field-level</code></p>')

    w('<div class="info">')
    w("<strong>Scope statement:</strong> This report summarizes a one-time exploratory analysis "
      "of field-boundary geometry, CDL crop data, and NASA POWER weather across three growers. "
      "Soil analysis was <strong>not</strong> requested and is <strong>not</strong> included.")
    w("</div>")

    # ---- Section 1: Dataset Scope ----
    w("<h2>1. Dataset Scope</h2>")
    w("<table>")
    for i, row in enumerate(GROWER_TABLE):
        tag = "th" if i == 0 or row[0] == "Total" else "td"
        w("<tr>" + "".join(f"<{tag}>{esc(c)}</{tag}>" for c in row) + "</tr>")
    w("</table>")
    w("<p>Each grower occupies a single farm. The runtime data root is "
      "<code>~/my-farm-advisor-runtime/data-pipeline/</code>. Growers were discovered "
      "dynamically from the <code>growers/</code> directory \u2014 no manual filter was applied.</p>")

    # ---- Section 2: Data Layers ----
    w("<h2>2. Data Layers</h2>")
    w("<table>")
    for i, row in enumerate(LAYERS_TABLE):
        tag = "th" if i == 0 else "td"
        w("<tr>" + "".join(f"<{tag}>{esc(c)}</{tag}>" for c in row) + "</tr>")
    w("</table>")

    # ---- Section 3: Comparison Levels ----
    w("<h2>3. Comparison Levels</h2>")
    w("<ul>")
    for label, desc in COMPARISON_LEVELS:
        w(f"<li><strong>{esc(label)}:</strong> {esc(desc)}</li>")
    w("</ul>")

    # ---- Sections 4-6: Categories ----
    section_num = 4
    for cat in CATEGORIES:
        w(f"<h2>{section_num}. {esc(cat['title'])}</h2>")

        # Category analysis block
        w(f'<div class="category-analysis"><strong>Analysis:</strong> {esc(cat["analysis"])}</div>')

        w(f"<h3>{section_num}.1 Outputs</h3>")

        # CSV files as a bullet list under a sub-heading
        if cat["csv_outputs"]:
            w("<h4>Data Files</h4>")
            w("<ul>")
            for csv_out in cat["csv_outputs"]:
                w(f'<li><code>{esc(csv_out["filename"])}</code> \u2014 {esc(csv_out["desc"])}</li>')
            w("</ul>")

        # PNG gallery - text before image, margin-bottom between bundles
        if cat["png_outputs"]:
            w('<div class="gallery">')
            for png_out in cat["png_outputs"]:
                is_per_grower = "template" in png_out
                if is_per_grower:
                    # Show first grower image, describe once
                    img_file = png_out["actual_files"][0]
                    w('<div class="figure">')
                    w(f'<div class="filename">{esc(png_out["template"])}</div>')
                    w(f'<div class="desc">{esc(png_out["desc"])}</div>')
                    w(f'<div class="analysis">{esc(png_out["analysis"])}</div>')
                    w(f'<img src="{esc(img_file)}" alt="{esc(png_out["template"])}">')
                    w("</div>")
                else:
                    w('<div class="figure">')
                    w(f'<div class="filename">{esc(png_out["filename"])}</div>')
                    w(f'<div class="desc">{esc(png_out["desc"])}</div>')
                    w(f'<div class="analysis">{esc(png_out["analysis"])}</div>')
                    w(f'<img src="{esc(png_out["filename"])}" alt="{esc(png_out["filename"])}">')
                    w("</div>")
            w("</div>")

        section_num += 1

    # ---- Key Observations ----
    w(f"<h2>{section_num}. Key Observations</h2>")
    w("<ul>")
    for label, desc in KEY_OBSERVATIONS:
        w(f"<li><strong>{esc(label)}:</strong> {esc(desc)}</li>")
    w("</ul>")
    section_num += 1

    # ---- Limitations ----
    w(f"<h2>{section_num}. Limitations</h2>")
    w('<div class="limitation">')
    w("<ul>")
    for label, desc in LIMITATIONS:
        w(f"<li><strong>{esc(label)}:</strong> {esc(desc)}</li>")
    w("</ul>")
    w("</div>")

    w('<p style="margin-top:3em; color:#888; font-size:0.85em;">')
    w("Report generated from <code>eda-field-level</code> outputs &mdash; no soil analysis performed.")
    w("</p>")

    w("</body>")
    w("</html>")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX generator
# ---------------------------------------------------------------------------

def _build_docx(output_dir: Path, grower_slugs: list[str], output_path: Path):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # --- Style setup ---
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

    # --- Title ---
    doc.add_heading("EDA Field-Level Analysis Report", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Generated: 2026-07-01  |  Subskill: my-farm-advisor/eda/eda-field-level")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    run = p.add_run("Scope statement: ")
    run.bold = True
    p.add_run(
        "This report summarizes a one-time exploratory analysis of field-boundary geometry, "
        "CDL crop data, and NASA POWER weather across three growers. Soil analysis was not "
        "requested and is not included."
    )

    # --- 1. Dataset Scope ---
    doc.add_heading("1. Dataset Scope", level=2)
    table = doc.add_table(rows=len(GROWER_TABLE), cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(GROWER_TABLE):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph(
        "Each grower occupies a single farm. The runtime data root is "
        "~/my-farm-advisor-runtime/data-pipeline/. Growers were discovered dynamically "
        "from the growers/ directory \u2014 no manual filter was applied."
    )

    # --- 2. Data Layers ---
    doc.add_heading("2. Data Layers", level=2)
    table = doc.add_table(rows=len(LAYERS_TABLE), cols=4)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(LAYERS_TABLE):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
            if i == 0:
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # --- 3. Comparison Levels ---
    doc.add_heading("3. Comparison Levels", level=2)
    for label, desc in COMPARISON_LEVELS:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(desc)

    # --- Sections 4-6: Categories ---
    section_num = 4
    for cat in CATEGORIES:
        doc.add_heading(f"{section_num}. {cat['title']}", level=2)

        # Category analysis
        p = doc.add_paragraph()
        run = p.add_run("Analysis: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB8, 0x5B, 0x14)
        p.add_run(cat["analysis"])

        doc.add_heading(f"{section_num}.1 Outputs", level=3)

        # CSV files
        if cat["csv_outputs"]:
            p = doc.add_paragraph()
            run = p.add_run("Data Files")
            run.bold = True
            for csv_out in cat["csv_outputs"]:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(csv_out["filename"])
                run.font.color.rgb = RGBColor(0x2A, 0x7A, 0x2A)
                p.add_run(f" \u2014 {csv_out['desc']}")

        # PNG outputs - text before image, spacing between bundles
        for png_out in cat["png_outputs"]:
            is_per_grower = "template" in png_out

            if is_per_grower:
                # Describe once, embed all images
                p = doc.add_paragraph()
                run = p.add_run(png_out["template"])
                run.bold = True
                p = doc.add_paragraph()
                run = p.add_run(f"Description: ")
                run.bold = True
                p.add_run(png_out["desc"])
                p = doc.add_paragraph()
                run = p.add_run("Analysis: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0xB8, 0x5B, 0x14)
                run.italic = True
                p.add_run(png_out["analysis"])
                for fname in png_out["actual_files"]:
                    img_path = output_dir / fname
                    if img_path.exists():
                        p = doc.add_paragraph()
                        run = p.add_run(fname)
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                        img_w = png_out.get("docx_img_width", 5.0)
                        doc.add_picture(str(img_path), width=Inches(img_w))

            else:
                p = doc.add_paragraph()
                run = p.add_run(png_out["filename"])
                run.bold = True
                p = doc.add_paragraph()
                run = p.add_run(f"Description: ")
                run.bold = True
                p.add_run(png_out["desc"])
                p = doc.add_paragraph()
                run = p.add_run("Analysis: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0xB8, 0x5B, 0x14)
                run.italic = True
                p.add_run(png_out["analysis"])
                img_path = output_dir / png_out["filename"]
                if img_path.exists():
                    doc.add_picture(str(img_path), width=Inches(5.0))

            # Blank line separator between bundles
            doc.add_paragraph()

        section_num += 1

    # --- Key Observations ---
    doc.add_heading(f"{section_num}. Key Observations", level=2)
    for label, desc in KEY_OBSERVATIONS:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(desc)
    section_num += 1

    # --- Limitations ---
    doc.add_heading(f"{section_num}. Limitations", level=2)
    for label, desc in LIMITATIONS:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(desc)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    output_dir = resolve_output_dir()
    grower_slugs = resolve_grower_slugs(output_dir)

    # Generate HTML
    html_content = _build_html(output_dir, grower_slugs)
    html_path = output_dir / "eda_field_level_report.html"
    html_path.write_text(html_content, encoding="utf-8")
    print(f"Saved: {html_path}")

    # Generate DOCX
    docx_path = output_dir / "eda_field_level_report.docx"
    _build_docx(output_dir, grower_slugs, docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    main()
